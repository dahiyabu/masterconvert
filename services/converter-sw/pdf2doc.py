import os
import shutil
import subprocess
import config
from init import logger
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
from pdf2docx import Converter
from docx2pdf import convert


# ✅ Detect if PDF is scanned
def is_scanned_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    for page in doc:
        if page.get_text("text").strip():
            return False  # Has text
    return True  # No text

# ✅ Convert scanned PDF to DOCX using OCR
def convert_scanned_pdf_to_docx(input_path, output_path):
    try:
        pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_PATH
        #pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Update this path
        #os.environ['TESSDATA_PREFIX'] = r'C:\Program Files\Tesseract-OCR\tessdata'+os.sep
        doc = Document()
        pdf_doc = fitz.open(input_path)

        for page_index in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_index)
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Use image_to_data to get structured text
            data = pytesseract.image_to_data(img, config="--psm 6", output_type=pytesseract.Output.DATAFRAME)
            data = data.dropna(subset=["text"])
            if data.empty:
                continue

            #doc.add_paragraph(f"--- Page {page_index + 1} ---")

            # Group by line and sort left-to-right
            for _, line in data.groupby("line_num"):
                line_text = " ".join(line.sort_values("left")["text"])
                if line_text.strip():
                    doc.add_paragraph(line_text.strip())

        doc.save(output_path)
        logger.info(f"✅ OCR-based DOCX saved at {output_path}")
        return True

    except Exception as e:
        logger.error(f"❌ OCR conversion failed: {e}")
        return False

def convert_from_pdf(input_path, output_path, dest_format):
    if dest_format == 'docx':
        return convert_pdf_to_docx(input_path,output_path)
    elif dest_format == 'odt':
        temp_docx = output_path.replace(".odt", ".docx")
        if convert_pdf_to_docx(input_path,temp_docx):
            return convert_docx_to_odt(temp_docx,output_path)
    return False

# ✅ Main function to convert PDF to DOCX (smart fallback)
def convert_pdf_to_docx(input_path,output_path):
    try:
        if is_scanned_pdf(input_path):
            logger.info("📄 Detected scanned PDF – using OCR method for DOCX.")
            return convert_scanned_pdf_to_docx(input_path, output_path)

        cv = Converter(input_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        logger.info(f"✅ PDF converted to DOCX at: {output_path}")
        return True
    except Exception as e:
        logger.error(f"❌ Error: {e}")
        return False
    
def convert_docx_to_odt(docx_path, odt_path):
    try:
        output_dir = os.path.dirname(odt_path)
        subprocess.run([
            #r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            config.LIBREOFFICE_PATH,
            '--headless',
            '--convert-to', 'odt',
            '--outdir', output_dir,
            docx_path
        ], check=True)
        print(f"✅ DOCX converted to ODT at {odt_path}")
        return True
    except subprocess.CalledProcessError as e:
        print(f"❌ LibreOffice error: {e}")
        return False